#! /usr/bin/env python3
# -*- coding: utf-8 -*-
# @Author  : muyan.wang@centurygame.com
# @Date    : 2022/4/29
# @Desc    :
import re

from docx import Document
from collections import defaultdict
import yaml

with open("config.yaml", 'r', encoding='utf-8') as yamlFile:
    config_data = yaml.load(yamlFile, Loader=yaml.FullLoader)


def find_index(content: str, tag):
    if re.findall("^（(\d{0,9})）|^(\d{0,9})[．]", content) and any(re.findall("^（(\d{0,9})）|^(\d{0,9})[．]", content)[0]):
        a = re.findall("^（(\d{0,9})）|^(\d{0,9})[．]", content)[0][0] or re.findall("^（(\d{0,9})）|^(\d{0,9})[．]", content)[0][1]
        # print(re.split(r"([/(（])?(\d+)([\)）\.．])", content))
        b = re.split(r"([/(（])?(\d+)([\)）．])", content)[4]
        return f"{a}【{tag}】{b}"
    else:
        return content

def add_tag(text, da_data, one_text, three_text, a):
    key_name = f"{one_text}_{three_text}_{a}"
    ddd = da_data.get(key_name) or [""]
    text.insert_paragraph_before(text=f'【答案】{ddd[0]}')
    for index in ddd[1:]:
        text.insert_paragraph_before(text=index)
    text.insert_paragraph_before(text='【解析】')
    # tag = re.sub(r"[一二三四五六七八九十\d．、]", "", three_text)
    # text.insert_paragraph_before(text=f'【标签】{tag}')


def get_daan(file_name):
    document = Document(f"da_docx/{file_name}")
    import re

    da_data = {}

    one_index = 0

    three_index = 0

    for text in document.paragraphs:
        if text.text:
            if re.findall(config_data['big_title'], text.text):
                one_index = re.findall(config_data['big_title'], text.text)[0]
            elif re.match(config_data['填空题']['biaotiguanjianzi'], text.text) or re.match(config_data['简答题']['biaotiguanjianzi'], text.text) or re.match(config_data['判断题']['biaotiguanjianzi'], text.text) or re.match(config_data['多项选择题']['biaotiguanjianzi'], text.text) or re.match(config_data['单项选择题']['biaotiguanjianzi'], text.text):
                three_index = text.text
            else:
                if one_index and three_index:
                    data = re.split(r"([（])?(\d+)([）．])", text.text)
                    if not data[0]:
                        data = data[2:]
                        key_biaoti = data[::4]
                        key_text = data[2::4]
                        for key, value in zip(key_biaoti, key_text):
                            key_name = f"{one_index}_{three_index}_{key}"
                            da_data.setdefault(key_name, []).append(value)
                    else:
                        da_data.setdefault(key_name, []).append(text.text)
    return da_data


def new_docx(file_name):
    da_data = get_daan(file_name)

    document = Document(f"old_docx/{file_name}")

    book_tag = ""
    one_text = ""
    three_text = ""
    before_a = 1

    for index, text in enumerate(document.paragraphs):
        if text.text:

            if book_tag and re.findall("^（\d{0,9}）|^(\d{0,9})[．]", text.text):
                a = re.findall("^（(\d{0,9})）|^(\d{0,9})[．]", text.text)[0][0] or re.findall("^（(\d{0,9})）|^(\d{0,9})[．]", text.text)[0][1]
                if a == "1":
                    text.insert_paragraph_before(text=config_data[book_tag]['shuoming'])
                else:
                    if before_a == int(a) - 1:
                        add_tag(text, da_data, one_text, three_text, int(a)-1)
                        before_a = int(a)
                    else:
                        # print("最后1题")
                        pass

            if re.match(config_data['big_title'], text.text):
                if before_a > 1:
                    add_tag(text, da_data, one_text, three_text, before_a)
                    before_a = 1
                one_text = re.findall(config_data['big_title'], text.text)[0]

            elif re.match(config_data['填空题']['biaotiguanjianzi'], text.text):
                if before_a > 1:
                    add_tag(text, da_data, one_text, three_text, before_a)
                    before_a = 1
                content = text.text
                p = text.clear()
                p.add_run("【题组】" + content)
                book_tag = "填空题"
                three_text = content

            elif re.match(config_data['简答题']['biaotiguanjianzi'], text.text):
                if before_a > 1:
                    add_tag(text, da_data, one_text, three_text, before_a)
                    before_a = 1
                content = text.text
                p = text.clear()
                p.add_run("【题组】" + content)
                book_tag = "简答题"
                three_text = content

            elif re.match(config_data['判断题']['biaotiguanjianzi'], text.text):
                if before_a > 1:
                    add_tag(text, da_data, one_text, three_text, before_a)
                    before_a = 1
                content = text.text
                p = text.clear()
                p.add_run("【题组】" + content)
                book_tag = "判断题"
                three_text = content

            elif re.match(config_data['多项选择题']['biaotiguanjianzi'], text.text):
                if before_a > 1:
                    add_tag(text, da_data, one_text, three_text, before_a)
                    before_a = 1
                content = text.text
                p = text.clear()
                p.add_run("【题组】" + content)
                book_tag = "多项选择题"
                three_text = content

            elif re.match(config_data['单项选择题']['biaotiguanjianzi'], text.text):
                if before_a > 1:
                    add_tag(text, da_data, one_text, three_text, before_a)
                    before_a = 1
                content = text.text
                p = text.clear()
                p.add_run("【题组】" + content)
                book_tag = "单项选择题"
                three_text = content

            elif book_tag == "填空题":
                content = re.sub("_+|\s+", "（）", text.text)
                new_content = find_index(content,  "填空题")
                p = text.clear()
                p.add_run(new_content)

            elif book_tag in ["简答题", "判断题", "多项选择题", "单项选择题"]:
                content = text.text
                new_content = find_index(content, book_tag)
                p = text.clear()
                p.add_run(new_content)


    document.save(f'new_docx/{file_name}')


from pathlib import Path

old_docx_path = Path("old_docx")
for file in old_docx_path.glob('**/*.docx'):
    print(f"当前执行的文件为=={file.name}, 文件报错的话，查看答案和源文件有什么不一致活着跳过当前文件")
    new_docx(file.name)
