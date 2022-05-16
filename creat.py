#! /usr/bin/env python3
# -*- coding: utf-8 -*-
# @Author  : muyan.wang@centurygame.com
# @Date    : 2022/4/29
# @Desc    :
import re

from docx import Document
from collections import defaultdict
import yaml

with open("config.yaml", 'r', encoding='utf-8') as yamlFile:
    config_data = yaml.load(yamlFile, Loader=yaml.FullLoader)


class Book:
    def __init__(self):
        self.max_title: str = ""
        self.subject_type: str = ""
        self.subject_index: str = ""


def max_title_bool(content: str):
    return re.match(config_data['big_title'], content)


def find_index(content: str, index, tag):
    if content.find(str(index)) == 0:
        return f"{content[:2]}【{tag}】{content[2:]}"
    elif content.find(str(index)) == 1:
        return f"{content[:3]}【{tag}】{content[3:]}"
    else:
        print(f"无法识别{content}首位")


def write_docx(new_testList: list, book: dict):
    if new_testList:
        new_document = Document()
        for index in new_testList:
            new_document.add_paragraph(index)
        new_document.save(f'new_docx/{book["max_title"]}.docx')
        new_testList.clear()
        book.clear()


document = Document("old_docx/460218312 《Python编程案例教程（第2版）》（摘课后习题）2022-4-24.docx")

print('节的数量：', len(document.sections))
print("檔案內含段落數：", len(document.paragraphs), "\n")

testList = [
    "【书ID】",
    "【标题】课后练习",
    "【章】",
    "【节】",
]
book = {}
new_testList = []
for text in document.paragraphs:
    if text.text:
        print(text.text)
        print("xxxxxxxxxxxxx")

        if new_testList and new_testList[-1][0].isupper() and not text.text[0].isupper():
            new_testList.append("答案：")
            new_testList.append("解析：")
            new_testList.append(f"标签：{book['tag']}")

        if text.style.name == "Heading 1" and max_title_bool(text.text):
            write_docx(new_testList, book)
            book["max_title"] = text.text
            new_testList.extend(testList)

        elif re.match(config_data['填空题']['biaotiguanjianzi'], text.text):
            new_testList.append(f"【题组】{text.text}")
            new_testList.append(config_data['填空题']['shuoming'])
            book['tag'] = '填空题'
            book['index'] = 1

        elif re.match(config_data['判断题']['biaotiguanjianzi'], text.text):
            new_testList.append(f"【题组】{text.text}")
            new_testList.append(config_data['判断题']['shuoming'])
            book['tag'] = '判断题'
            book['index'] = 1

        elif re.match(config_data['多项选择题']['biaotiguanjianzi'], text.text):
            new_testList.append(f"【题组】{text.text}")
            new_testList.append(config_data['多项选择题']['shuoming'])
            book['tag'] = '多项选择题'
            book['index'] = 1

        elif re.match(config_data['简答题']['biaotiguanjianzi'], text.text):
            new_testList.append(f"【题组】{text.text}")
            new_testList.append(config_data['简答题']['shuoming'])
            book['tag'] = '简答题'
            book['index'] = 1

        elif book.get('tag') == "填空题":
            content = re.sub("_+|\s+", "（ ）", text.text)
            if content:
                new_content = find_index(content, book['index'], book['tag'])
                if new_content:
                    new_testList.append(new_content)
                    new_testList.append("答案：")
                    new_testList.append("解析：")
                    new_testList.append("标签：")
                    book['index'] += 1

        elif book.get('tag') == "判断题":
            content = text.text.replace("（    ）", "")
            if content:
                new_content = find_index(content, book['index'], book['tag'])
                if new_content:
                    new_testList.append(new_content)
                    new_testList.append("答案：")
                    new_testList.append("解析：")
                    new_testList.append("标签：")
                    book['index'] += 1

        elif book.get('tag') == "简答题":
            content = text.text
            if content:
                new_content = find_index(content, book['index'], book['tag'])
                if new_content:
                    new_testList.append(new_content)
                    new_testList.append("答案：")
                    new_testList.append("解析：")
                    new_testList.append("标签：")
                    book['index'] += 1

        elif book.get('tag') == "多项选择题":
            content = text.text
            if content:
                if text.text[0].isupper() and text.text[1] == "．":
                    xz_list = text.text.split("\t")
                    for xzt in xz_list:
                        if xzt:
                            new_testList.append(xzt)
                else:
                    new_content = find_index(content, book['index'], book['tag'])
                    if new_content:
                        new_testList.append(new_content)
                        book['index'] += 1

write_docx(new_testList, book)  # 写最后一章节
